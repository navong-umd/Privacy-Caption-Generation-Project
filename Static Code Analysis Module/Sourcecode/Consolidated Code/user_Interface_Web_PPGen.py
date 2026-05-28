# --- requirements (run once) ---
# pip install python-docx

# --- at top of app.py (imports) ---
from flask import (
    Flask, request, redirect, url_for, render_template_string,
    send_from_directory, send_file, flash, session
)
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pathlib import Path
from datetime import date, datetime
import pandas as pd
import os
# --- paths: put below CSV_OUTPUT_FOLDER definition ---
TEMPLATE_DIR = Path(
    r"C:\Users\2015n\OneDrive\Desktop\Privacy Policy Generation Project\Static Code Analysis Module\template"
)
TEMPLATE_FILE = TEMPLATE_DIR / "Sample Privacy Policy Template.docx"

CSV_OUTPUT_FOLDER = Path(
    r"C:\Users\2015n\OneDrive\Desktop\Privacy Policy Generation Project\Static Code Analysis Module\output_csv"
)

POLICIES_OUTPUT_DIR = CSV_OUTPUT_FOLDER / "policies"
POLICIES_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def _safe(val):
    return "" if val is None else str(val)

# ====== Session answer helpers ======
def answers():
    if "answers" not in session:
        session["answers"] = {}
    return session["answers"]

def _load_master_and_sections():
    """
    Reads the two Excel files we already write on Publish and returns:
      - master (dict) from questionnaire_master.xlsx first row
      - by_sections (dict[str, DataFrame]) for reference if needed
    """
    master_path = CSV_OUTPUT_FOLDER / "questionnaire_master.xlsx"
    by_path = CSV_OUTPUT_FOLDER / "questionnaire_by_section.xlsx"

    if not master_path.exists():
        raise FileNotFoundError(f"Missing {master_path}")
    if not by_path.exists():
        raise FileNotFoundError(f"Missing {by_path}")

    master_df = pd.read_excel(master_path)
    if master_df.empty:
        raise ValueError("questionnaire_master.xlsx has no rows")
    master = {c: _safe(master_df.iloc[0].get(c, "")) for c in master_df.columns}

    # Read all sheets (optional, but we return them in case you want to expand)
    by_sections = {}
    with pd.ExcelFile(by_path) as xf:
        for sheet in xf.sheet_names:
            by_sections[sheet] = pd.read_excel(xf, sheet_name=sheet)

    return master, by_sections

def _build_template_context(master: dict) -> dict:
    """
    Build a flat context mapping from master row suitable for {{placeholders}} in the .docx template.
    Feel free to add more keys if your template needs them.
    """
    # resolve policy title
    ttl = answers().get("Final - Title", {})
    ptitle = ttl.get("ptitle_custom") if (ttl.get("ptitle_opt") == "custom" and ttl.get("ptitle_custom")) else (ttl.get("ptitle_opt") or "Privacy Policy")

    # version date (already set in master)
    eff_date = _safe(master.get("last_update_date") or master.get("submission_date") or "")

    # Booleans come in as "Yes"/"No" strings in master: keep them as-is for readability.
    ctx = {
        # Header
        "policy_title": ptitle,
        "effective_date": eff_date,

        # Organization / contact
        "org_name": _safe(master.get("org_name")),
        "app_name": _safe(master.get("app_name")),
        "website_url": _safe(master.get("website_url")),
        "contact_email": _safe(master.get("contact_email")),
        "contact_address": _safe(master.get("contact_address")),
        "additional_contacts": _safe(master.get("additional_contacts")),
        "phone": _safe(master.get("phone") or ""),

        # Collection
        "collect_personal_data": _safe(master.get("collect_personal_data")),
        "data_categories": _safe(master.get("data_categories")),
        "collection_methods": _safe(master.get("collection_methods")),
        "source_third_party": _safe(master.get("source_third_party")),

        # Additional Collection
        "device_collected": _safe(master.get("device_data_collected")),
        "device_categories": _safe(master.get("device_data_categories")),
        
        # Purposes
        "purpose_core_functions": _safe(master.get("purpose_core_functions")),
        "purpose_analytics": _safe(master.get("purpose_analytics")),
        "purpose_marketing": _safe(master.get("purpose_marketing")),
        "purpose_personalization": _safe(master.get("purpose_personalization")),
        "purpose_compliance": _safe(master.get("purpose_compliance")),

        # Tracking / analytics
        "cookies": _safe(master.get("cookies")),
        "maps_apis": _safe(master.get("maps_apis")),
        "ga_used": _safe(master.get("ga_used")),
        "ga_feature": _safe(master.get("ga_feature")),

        # Sharing
        "share_advertisers": _safe(master.get("share_advertisers")),
        "share_affiliates": _safe(master.get("share_affiliates")),
        "share_partners": _safe(master.get("share_partners")),

        # Security & retention
        "security_measures": _safe(master.get("security_measures")),
        "retention_policy": _safe(master.get("retention_policy")),
        "breach_notification": _safe(master.get("breach_notification")),

        # Rights
        "rights_access": _safe(master.get("rights_access")),
        "rights_correct": _safe(master.get("rights_correct")),
        "rights_delete": _safe(master.get("rights_delete")),
        "rights_portability": _safe(master.get("rights_portability")),
        "rights_object": _safe(master.get("rights_object")),
        "rights_withdraw_consent": _safe(master.get("rights_withdraw_consent")),
        "rights_request_process": _safe(master.get("rights_request_process")),

        # Children & transfers & CCPA
        "children_under_13": _safe(master.get("children_under_13")),
        "parental_consent": _safe(master.get("parental_consent")),
        "transfer_outside_region": _safe(master.get("transfer_outside_region")),
        "transfer_mechanisms": _safe(master.get("transfer_mechanisms")),
        "ccpa_sale_share": _safe(master.get("ccpa_sale_share")),
        "opt_out_links": _safe(master.get("opt_out_links")),

        # Misc
        "notes_internal": _safe(master.get("notes_internal")),
        "submission_id": _safe(master.get("submission_id")),
        "submitted_at": _safe(master.get("submitted_at")),
    }
    return ctx

def _replace_all(text: str, ctx: dict) -> str:
    # naive {{key}} replacement
    out = text
    for k, v in ctx.items():
        out = out.replace("{{" + k + "}}", str(v))
    return out

def _replace_placeholders(document: Document, ctx: dict):
    # paragraphs
    for p in document.paragraphs:
        for r in p.runs:
            r.text = _replace_all(r.text, ctx)
    # tables
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = _replace_all(r.text, ctx)

    # Headers / footers (often forgotten)
    for sec in document.sections:
        for p in sec.header.paragraphs:
            for r in p.runs:
                r.text = _replace_all(r.text, ctx)
        for p in sec.footer.paragraphs:
            for r in p.runs:
                r.text = _replace_all(r.text, ctx)

def generate_policy_docx() -> Path:
    """
    Opens the template, replaces {{placeholders}}, saves a new docx under POLICIES_OUTPUT_DIR.
    Returns the path to the generated file.
    """
    master, _ = _load_master_and_sections()
    ctx = _build_template_context(master)

    # Load template or fall back to a blank document
    try:
        if TEMPLATE_FILE.exists():
            doc = Document(str(TEMPLATE_FILE))
        else:
            # Fallback minimal structure if the template is missing
            doc = Document()
            doc.add_heading(ctx.get("policy_title") or "Privacy Policy", level=1)
            doc.add_paragraph(f"Effective Date: {ctx.get('effective_date','')}")
            doc.add_paragraph("")  # spacer
            # You can add more fallback structure here if needed
    except (PackageNotFoundError, ValueError) as e:
        # Corrupt template fallback
        doc = Document()
        doc.add_heading(ctx.get("policy_title") or "Privacy Policy", level=1)
        doc.add_paragraph(f"(Template could not be opened: {e})")
        doc.add_paragraph(f"Effective Date: {ctx.get('effective_date','')}")

    # Perform replacements on whatever document we have
    _replace_placeholders(doc, ctx)

    # Filename based on app/org and timestamp
    base = (ctx.get("app_name") or ctx.get("org_name") or "PrivacyPolicy").strip().replace(" ", "_")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"{base}_Privacy_Policy_{stamp}.docx"
    out_path = POLICIES_OUTPUT_DIR / out_name
    doc.save(out_path)

    # store filename only (not full path!)
    session["latest_policy_fname"] = out_path.name
    session.modified = True

    return out_path